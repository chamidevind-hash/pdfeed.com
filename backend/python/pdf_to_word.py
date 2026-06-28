import argparse
import sys


def convert_with_pdf2docx(source, destination):
    from pdf2docx import Converter

    converter = Converter(source)
    try:
        converter.convert(destination)
    finally:
        converter.close()


def add_text_run(paragraph, span):
    text = span.get("text", "")
    if not text:
        return

    run = paragraph.add_run(text)
    size = span.get("size")
    if size:
        from docx.shared import Pt

        run.font.size = Pt(max(1, min(float(size), 96)))

    font_name = span.get("font")
    if font_name:
        run.font.name = font_name

    flags = int(span.get("flags", 0))
    run.bold = bool(flags & 16)
    run.italic = bool(flags & 2)


def convert_with_fallback(source, destination):
    import fitz
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Inches

    pdf = fitz.open(source)
    document = Document()

    try:
        for page_index, page in enumerate(pdf):
            page_width = max(float(page.rect.width), 1)
            page_height = max(float(page.rect.height), 1)
            section = document.sections[-1]
            section.page_width = Inches(page_width / 72)
            section.page_height = Inches(page_height / 72)

            content = page.get_text("dict", sort=True)
            for block in content.get("blocks", []):
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        paragraph = document.add_paragraph()
                        for span in line.get("spans", []):
                            add_text_run(paragraph, span)
                elif block.get("type") == 1 and block.get("image"):
                    image = block["image"]
                    try:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(
                            _bytes_stream(image),
                            width=Inches(min(page_width / 72, 7.5)),
                        )
                    except Exception:
                        pass

            if page_index < len(pdf) - 1:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        document.save(destination)
    finally:
        pdf.close()


def _bytes_stream(data):
    from io import BytesIO

    return BytesIO(data)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--mode", choices=("pdf2docx", "fallback"), required=True)
    parser.add_argument("source")
    parser.add_argument("destination")
    args = parser.parse_args()

    if args.mode == "pdf2docx":
        convert_with_pdf2docx(args.source, args.destination)
    else:
        convert_with_fallback(args.source, args.destination)


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        print(str(error), file=sys.stderr)
        sys.exit(1)
